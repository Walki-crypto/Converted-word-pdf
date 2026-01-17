# src/converter.py
"""Utility functions to convert .docx files to PDF using ReportLab.
Pure Python solution with no external native dependencies required.
"""

import os
import tempfile
from pathlib import Path
from typing import Union

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

def convert_docx_to_pdf(input_path: Union[str, Path], output_path: Union[str, Path] = None) -> Path:
    """Convert a .docx file to PDF using ReportLab.
    Handles headings, paragraphs with basic styles, tables, and images.
    """
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    if input_path.suffix.lower() != ".docx":
        raise ValueError("Input file must have a .docx extension")

    if output_path is None:
        output_path = input_path.with_suffix('.pdf')
    else:
        output_path = Path(output_path)

    doc = Document(input_path)
    styles = getSampleStyleSheet()
    story = []

    # 1. Process Paragraphs and Headings
    for para in doc.paragraphs:
        style_name = para.style.name
        # Map Word styles to ReportLab styles
        rl_style = styles['Normal']
        if style_name.startswith('Heading'):
            level = style_name.replace('Heading', '').strip()
            if level == '1': rl_style = styles['Heading1']
            elif level == '2': rl_style = styles['Heading2']
            elif level == '3': rl_style = styles['Heading3']
            else: rl_style = styles['Heading1']

        # Build markup for bold, italic, underline
        markup = ""
        for run in para.runs:
            text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            markup += text
        
        if markup.strip() or style_name.startswith('Heading'):
            story.append(Paragraph(markup, rl_style))
            story.append(Spacer(1, 12))

    # 2. Process Tables
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        if data:
            tbl = Table(data, hAlign='LEFT')
            tbl.setStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ])
            story.append(tbl)
            story.append(Spacer(1, 12))

    # 3. Process Images
    for i, shape in enumerate(doc.inline_shapes):
        try:
            # Extract image bytes
            rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = doc.part.related_parts[rId]
            img_data = image_part.blob
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                tmp.write(img_data)
                tmp_path = tmp.name
            
            # Add to Story
            # Scale image if it's too wide
            available_width = LETTER[0] - 100 # Approx margins
            story.append(RLImage(tmp_path, width=min(250, available_width), preserveAspectRatio=True))
            story.append(Spacer(1, 12))
        except Exception:
            # Skip images that fail to process
            continue

    # Build PDF
    doc_pdf = SimpleDocTemplate(str(output_path), pagesize=LETTER)
    doc_pdf.build(story)
    return output_path

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python -m src.converter <input.docx> [output.pdf]")
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    try:
        res = convert_docx_to_pdf(inp, out)
        print(f"PDF generated: {res}")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
