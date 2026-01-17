# create_sample.py
"""Genera un archivo .docx de ejemplo para probar el convertidor.
Incluye:
- Título (Heading 1)
- Párrafo con estilos (negrita, cursiva)
- Tabla simple
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt


def build_sample_docx(path: Path):
    doc = Document()
    # Título
    heading = doc.add_heading('Ejemplo de Conversión', level=1)
    # Párrafo con estilos
    p = doc.add_paragraph()
    p.add_run('Este es un párrafo de ejemplo con ').bold = True
    p.add_run('negrita, ').italic = True
    p.add_run('cursiva y texto normal.')
    # Tabla
    table = doc.add_table(rows=2, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Columna 1'
    hdr_cells[1].text = 'Columna 2'
    hdr_cells[2].text = 'Columna 3'
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Dato A1'
    row_cells[1].text = 'Dato B1'
    row_cells[2].text = 'Dato C1'
    # Guardar
    doc.save(path)

if __name__ == "__main__":
    out_path = Path(__file__).parent / "sample.docx"
    build_sample_docx(out_path)
    print(f"Archivo de muestra creado en: {out_path}")
